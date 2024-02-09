import tkinter as tk
from collections import defaultdict
from tkinter import messagebox
from datetime import date, datetime
from tkinter.filedialog import askdirectory
from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from tkinter import simpledialog
from docx.shared import RGBColor
from tkcalendar import Calendar
from tkcalendar import DateEntry

import os
import re
import math
import json
import tempfile

class App(tk.Tk):

    BASE_COST = 89.5  # базовая стоимость бандероли
    STEP_COST = 3.5   # стоимость за шаг в 20 грамм
    LETTER_COST = 29.0  # стоимость письма простого
    REGISTERED_LETTER_COST = 67.0 # стоимость письма заказного
    NDS = 1.2 # Ндс 20%
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.title("Отчет по почте")

        self.load_settings_from_file() # загружаем настройки при старте приложения

        self.total_weight = 0
        self.total_cost = 0.0
        self.total_parcels = 0

        self.weights = []

        self.numbers_entered = []  # Список для хранения введенных значений простых писем
        self.numbers_entered_reg = []  # Список для хранения введенных значений заказных писем

        # Получить путь к рабочему столу пользователя
        self.desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
        
        self.create_widgets()

# Центруем виджет по середине
    def center_window(self, width, height):
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        x = (screen_width // 2) - (width // 2)
        y = (screen_height // 2) - (height // 2)

        self.geometry(f"{width}x{height}+{x}+{y}")

# Это главное окно с последующими кнопками
    def create_widgets(self):
        self.geometry("600x330")  # Прямо здесь задаем размер корневого окна
 
        # Кнопка подсчета за день
        self.calculate_button = tk.Button(
            self, text="Подсчет бандеролей", command=self.open_weight_window
        )
        self.calculate_button.pack(pady=10)

        # Кнопка подсчета писем (простых)
        self.letters2_button = tk.Button(
            self, text="Подсчет заказных писем", command=self.calculate_registered_letters
        )
        self.letters2_button.pack(pady=10)

        # Кнопка подсчета писем (заказных)
        self.letters_button = tk.Button(
            self, text="Подсчет простых писем", command=self.calculate_letters
        )
        self.letters_button.pack(pady=10)

        # Кнопка посылок
        self.parcels_button = tk.Button(
            self, text="Подсчет посылок", command=self.open_parcels_window
        )
        self.parcels_button.pack(pady=10)

        # Кнопка подсчета за месяц
        self.monthly_button = tk.Button(
            self, text="Отчет за месяц", command=self.ask_month_input
        )
        self.monthly_button.pack(pady=10)

        self.doc_button = tk.Button(
            self, text="Создать обложку", command=self.open_cover_window
        )
        self.doc_button.pack(pady=10)

        # Надпись "by.Borzzz" в нижнем правом углу
        by_label = tk.Label(self, text="by.Borzzz", fg="gray")
        by_label.pack(side=tk.RIGHT, padx=10, pady=10)

        # Кнопка настроек
        self.settings_button = tk.Button(
            self, text="⚙️", command=self.open_settings_window, font=("default", 12)
        )
        self.settings_button.pack(side=tk.BOTTOM, anchor=tk.SW, padx=10, pady=10)

        self.center_window(600, 330)

# Открывается кнопка ввода веса бандеролей в день        
    def open_weight_window(self):
        self.weight_window = tk.Toplevel(self)
        self.weight_window.title("Подсчет веса бандеролей")
        
        self.weight_window.geometry("350x500")
        
        #Выравнивание посередине
        screen_width = self.weight_window.winfo_screenwidth()
        screen_height = self.weight_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 350) // 2 - 500
        y_coordinate = (screen_height - 500) // 2

        # Устанавливаем положение окна по центру
        self.weight_window.geometry(f"350x500+{x_coordinate}+{y_coordinate}")
        
        
        self.weight_label = tk.Label(self.weight_window, text="Введите вес бандероли (в граммах):")
        self.weight_label.pack(pady=10)
        
        self.weight_entry = tk.Entry(self.weight_window)
        self.weight_entry.pack(pady=10)
        self.weight_entry.focus_set()

        by_label = tk.Label(self.weight_window, text="by.Borzzz", fg="gray")
        by_label.pack(side=tk.BOTTOM, anchor=tk.SE, padx=10, pady=10)

# Создание Listbox для отображения введенных весов
        self.weights_listbox_label = tk.Label(self.weight_window, text="Список введённых значений:")
        self.weights_listbox_label.pack()
        self.weights_listbox = tk.Listbox(self.weight_window)
        self.weights_listbox.pack()
        
        self.delete_selected_weight_button = tk.Button(
            self.weight_window, text="Удалить выбранный вес", command=self.delete_selected_weight
        )
        self.delete_selected_weight_button.pack(pady=10)
        
        self.finish_button = tk.Button(
            self.weight_window, text="Закончить подсчет", command=self.finish_weight_calculation
        )
        self.finish_button.pack(pady=10)
        
        self.weight_entry.bind("<Return>", self.add_weight)

# Это округление введенных бандеролей до целого четного числа равному 20
    def round_weight(self, weight):
        return math.ceil(weight / 20.0) * 20

# Добавление бандеролей в общий список до подсчета
    def add_weight(self, event=None):
        try:
            weight = float(self.weight_entry.get())
            # Округляем вес
            rounded_weight = self.round_weight(weight)
            if rounded_weight < 120 or rounded_weight > 2000:
                raise ValueError("Введите валидный вес (от 120 до 2000).")
            self.weights.append(rounded_weight)  # Добавляем округленный вес
            self.total_weight += rounded_weight  # Используем округленный вес для общего веса
            self.total_parcels += 1
            self.total_cost += self.calculate_cost(rounded_weight)  # Рассчитываем стоимость по округленному весу
            self.weights_listbox.insert(tk.END, f"{rounded_weight} грамм")
            self.weight_entry.delete(0, tk.END)
        except ValueError as e:
            messagebox.showerror("Ошибка", str(e))
        finally:
            self.weight_entry.focus_set()

# Удаление выбранного веса из списка   
    def delete_last_weight(self):
        if not self.weights:
            messagebox.showinfo("Информация", "Список весов пуст.")
        else:
            last_weight = self.weights.pop()
            self.total_weight -= last_weight
            self.total_parcels -= 1
            self.total_cost -= self.calculate_cost(last_weight)
            self.weights_listbox.delete(tk.END)
        
        # Сообщение и фокусировка после изменения
        self.weight_entry.focus_set()

    def finish_weight_calculation(self):
        self.weight_window.destroy()
        self.open_date_window()

# Расчет стоимости бандеролей, идет по настройкам с выставленными значениями.
    def calculate_cost(self, weight):
        additional_cost = max(0, (weight - 120) // 20 * self.STEP_COST)
        return self.BASE_COST + additional_cost

# Открывается поле ввода даты перед сохранением списка с бандеролями       
    def open_date_window(self):
        self.date_window = tk.Toplevel(self)
        self.date_window.title("Выберите дату")
       
        window_width = 300
        window_height = 300

        # Получаем размеры экрана
        screen_width = self.date_window.winfo_screenwidth()
        screen_height = self.date_window.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # Устанавливаем положение окна
        self.date_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # Создаем календарь
        self.cal = Calendar(self.date_window, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        self.cal.pack(pady=10)

        # Функция сохранения результатов при выборе даты
        def save_results():
            selected_date = self.cal.get_date()
            self.save_results(selected_date)

        # Кнопка для сохранения результатов
        self.save_button = tk.Button(
            self.date_window, text="Сформировать список", command=save_results
        )
        self.save_button.pack(pady=10)

# Сохранение результата списка бандеролей с датой
    def save_results(self, event=None):
        selected_date = self.cal.get_date()
        try:
            current_date = datetime.strptime(selected_date, "%d.%m.%Y").date()
            
            result_string = (f"Итого за {current_date.strftime('%d.%m.%Y')} отправлено {self.total_parcels}"
                             f" бандеролей весом {self.total_weight:.2f} грамм на сумму {self.total_cost:.2f} рублей.\n")
            
            custom_path = self.custom_path
            filename = os.path.join(custom_path, f"Списки бандеролей.txt")

            with open(filename, "a", encoding='utf-8') as file:
                file.write(result_string)
        
            messagebox.showinfo("Успешно", "Результаты сохранены.")
            self.date_window.destroy()
        
            self.weights.clear()
            self.total_weight = 0
            self.total_cost = 0.0
            self.total_parcels = 0
        except ValueError:
            messagebox.showerror("Ошибка", "Введите дату в правильном формате (дд.мм.гггг).")

# Отыкрывается кнопка ЗАКАЗНЫХ писем  
    def calculate_registered_letters(self):
        # Открытие нового окна для ввода количества писем и даты
        self.letters_window = tk.Toplevel()
        self.letters_window.title("Подсчет заказных писем")
        self.letters_window.geometry("350x500")

        screen_width = self.letters_window.winfo_screenwidth()
        screen_height = self.letters_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 350) // 2- 500
        y_coordinate = (screen_height - 500) // 2

        # Устанавливаем положение окна по центру
        self.letters_window.geometry(f"350x500+{x_coordinate}+{y_coordinate}")

        self.numbers_entered_reg = []  # Список для хранения введенных значений
        
        # Ввод количества писем
        self.quantity_label = tk.Label(self.letters_window, text="Введите количество писем и нажмите Enter:")
        self.quantity_label.pack(pady=(20,5))
        
        self.quantity_entry = tk.Entry(self.letters_window)
        self.quantity_entry.pack(pady=5)
        self.quantity_entry.bind("<Return>", self.add_to_list_reg)  # Привязка к кнопке Enter
        self.quantity_entry.focus_set()
        
        # Список введенных значений
        self.listbox_label = tk.Label(self.letters_window, text="Список введённых значений:")
        self.listbox_label.pack(pady=(10,0))  # Отступ сверху перед надписью
        self.listbox = tk.Listbox(self.letters_window)
        self.listbox.pack(pady=(0,5))

        # Кнопка "удалить выбранное" для удаления конкретного введенного результата
        self.delete_selected_button_reg = tk.Button(self.letters_window, text="Удалить выбранное", command=self.remove_selected_reg)
        self.delete_selected_button_reg.pack(pady=(5,10))

        # Кнопка "Закончить подсчет" для открытия календаря
        self.finish_button = tk.Button(self.letters_window, text="Закончить подсчет", command=self.open_calendar_reg)
        self.finish_button.pack(pady=(5, 10))

        # Надпись "by.Borzzz" в нижнем правом углу
        by_label = tk.Label(self.letters_window, text="by.Borzzz", fg="gray")
        by_label.pack(side=tk.BOTTOM, anchor=tk.SE, padx=10, pady=10)

    def open_calendar_reg(self):
        # Создание календаря
        self.calendar_window = tk.Toplevel(self.letters_window)
        self.calendar_window.title("Выберите дату")
        self.calendar_window.geometry("300x300")

        screen_width = self.calendar_window.winfo_screenwidth()
        screen_height = self.calendar_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2
        y_coordinate = (screen_height - 300) // 2

        # Устанавливаем положение окна по центру
        self.calendar_window.geometry(f"300x300+{x_coordinate}+{y_coordinate}")

        self.cal = Calendar(self.calendar_window, selectmode="day", year=datetime.now().year,
                            month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        self.cal.pack(pady=10)

        # Кнопка для сохранения результатов
        self.save_button = tk.Button(self.calendar_window, text="Сохранить", command=self.calculate_and_save_result_reg)
        self.save_button.pack(pady=10)

# Это лист, где отображаются введенные письма (как в памяти так и в окне в виде списка)    
    def add_to_list_reg(self, event):
        # Попытка преобразовать введенные данные в число и добавление в список
        try:
            num_letters = int(self.quantity_entry.get())
            self.numbers_entered_reg.append(num_letters)  # Добавление числа в список
            self.listbox.insert(tk.END, num_letters)  # Вывод числа в интерфейсе
            self.quantity_entry.delete(0, tk.END)  # Очистка поля ввода
        except ValueError:
            tk.messagebox.showwarning("Ошибка", "Введите корректное число!")

# Если нажали кнопку удалить последний результат из списка писем.
    def remove_selected_reg(self):
        try:
            # Получить индекс выбранного элемента
            index = self.listbox.curselection()[0]
            # Удалить этот элемент из Listbox и из списка numbers_entered
            self.numbers_entered_reg.pop(index)
            self.listbox.delete(index)
        except IndexError:
            tk.messagebox.showwarning("Ошибка", "Выберите элемент для удаления")
        except Exception as e:
            tk.messagebox.showwarning("Ошибка", f"Произошла ошибка: {e}")
        
# Подсчет и сохранение итога по письмам
    def calculate_and_save_result_reg(self, event=None):
        # Ввод даты и подсчет итогаa
        selected_date = self.cal.get_date()
        selected_date = re.sub(r'[\s\\\/.,]', '.', selected_date)
        if not selected_date or not self.numbers_entered_reg:
            tk.messagebox.showwarning("Ошибка", "Введите все данные корректно!")
            return

        # Подсчет итога
        total = sum(self.numbers_entered_reg) * self.REGISTERED_LETTER_COST
        self.save_to_file_reg(total, sum(self.numbers_entered_reg), selected_date)

        # Отображение результата
        tk.messagebox.showinfo("Успешно", "Результаты сохранены.")
    
        # Закрытие окна ввода
        self.letters_window.destroy()

    def save_to_file_reg(self, total_result, total_registered_letters, date):
        custom_path = self.custom_path
        filename = os.path.join(custom_path, f"Списки заказных писем.txt")

        # Используем режим 'a' для добавления данных в конец файла
        with open(filename, 'a', encoding='utf-8') as file:
            file.write(f"Дата: {date} Количество писем: {total_registered_letters} Итого: {total_result} руб.\n")


# Отыкрывается кнопка ПРОСТЫХ писем  
    def calculate_letters(self):
        # Открытие нового окна для ввода количества писем и даты
        self.letters_window = tk.Toplevel()
        self.letters_window.title("Подсчет простых писем")
        self.letters_window.geometry("350x500")

        screen_width = self.letters_window.winfo_screenwidth()
        screen_height = self.letters_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 350) // 2- 500
        y_coordinate = (screen_height - 500) // 2

        # Устанавливаем положение окна по центру
        self.letters_window.geometry(f"350x500+{x_coordinate}+{y_coordinate}")

        self.numbers_entered = []  # Список для хранения введенных значений
        
        # Ввод количества писем
        self.quantity_label = tk.Label(self.letters_window, text="Введите количество писем и нажмите Enter:")
        self.quantity_label.pack(pady=(20,5))
        
        self.quantity_entry = tk.Entry(self.letters_window)
        self.quantity_entry.pack(pady=5)
        self.quantity_entry.bind("<Return>", self.add_to_list)  # Привязка к кнопке Enter
        self.quantity_entry.focus_set()
        
        # Список введенных значений
        self.listbox_label = tk.Label(self.letters_window, text="Список введённых значений:")
        self.listbox_label.pack(pady=(10,0))  # Отступ сверху перед надписью
        self.listbox = tk.Listbox(self.letters_window)
        self.listbox.pack(pady=(0,5))

        # Кнопка "удалить выбранное" для удаления конкретного введенного результата
        self.delete_selected_button = tk.Button(self.letters_window, text="Удалить выбранное", command=self.remove_selected)
        self.delete_selected_button.pack(pady=(5,10))

        # Кнопка "Закончить подсчет" для открытия календаря
        self.finish_button = tk.Button(self.letters_window, text="Закончить подсчет", command=self.open_calendar)
        self.finish_button.pack(pady=(5, 10))

        by_label = tk.Label(self.letters_window, text="by.Borzzz", fg="gray")
        by_label.pack(side=tk.BOTTOM, anchor=tk.SE, padx=10, pady=10)

    def open_calendar(self):
        # Создание календаря
        self.calendar_window = tk.Toplevel(self.letters_window)
        self.calendar_window.title("Выберите дату")
        self.calendar_window.geometry("300x300")

        screen_width = self.calendar_window.winfo_screenwidth()
        screen_height = self.calendar_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2
        y_coordinate = (screen_height - 300) // 2

        # Устанавливаем положение окна по центру
        self.calendar_window.geometry(f"300x300+{x_coordinate}+{y_coordinate}")

        self.cal = Calendar(self.calendar_window, selectmode="day", year=datetime.now().year,
                            month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        self.cal.pack(pady=10)

        # Кнопка для сохранения результатов
        self.save_button = tk.Button(self.calendar_window, text="Сохранить", command=self.calculate_and_save_result)
        self.save_button.pack(pady=10)

# Это лист, где отображаются введенные письма (как в памяти так и в окне в виде списка)    
    def add_to_list(self, event):
        # Попытка преобразовать введенные данные в число и добавление в список
        try:
            num_letters = int(self.quantity_entry.get())
            self.numbers_entered.append(num_letters)  # Добавление числа в список
            self.listbox.insert(tk.END, num_letters)  # Вывод числа в интерфейсе
            self.quantity_entry.delete(0, tk.END)  # Очистка поля ввода
        except ValueError:
            tk.messagebox.showwarning("Ошибка", "Введите корректное число!")

# Если нажали кнопку удалить последний результат из списка писем.
    def remove_selected(self):
        try:
            # Получить индекс выбранного элемента
            index = self.listbox.curselection()[0]
            # Удалить этот элемент из Listbox и из списка numbers_entered
            self.numbers_entered.pop(index)
            self.listbox.delete(index)
        except IndexError:
            tk.messagebox.showwarning("Ошибка", "Выберите элемент для удаления")
        except Exception as e:
            tk.messagebox.showwarning("Ошибка", f"Произошла ошибка: {e}")

# Подсчет и сохранение итога по письмам
    def calculate_and_save_result(self, event=None):
        # Ввод даты и подсчет итога
        selected_date = self.cal.get_date()
        selected_date = re.sub(r'[\s\\\/.,]', '.', selected_date)
        if not selected_date or not self.numbers_entered:
            tk.messagebox.showwarning("Ошибка", "Введите все данные корректно!")
            return

        try:
            total = sum(self.numbers_entered) * self.LETTER_COST
            self.save_to_file(total, sum(self.numbers_entered), selected_date)

            tk.messagebox.showinfo("Успешно", "Результаты сохранены.")

            self.letters_window.destroy()
        except ValueError:
            tk.messagebox.showerror("Ошибка", "Ошибка при сохранении результатов.")

    def save_to_file(self, total_result, total_letters, date):
        custom_path = self.custom_path
        filename = os.path.join(custom_path, f"Списки простых писем.txt")

        # Используем режим 'a' для добавления данных в конец файла
        with open(filename, 'a', encoding='utf-8') as file:
            file.write(f"Дата: {date} Количество писем: {total_letters} Итого: {total_result} руб.\n")

#Функция удаления значений из листов
    def delete_selected_weight(self):
        selection = self.weights_listbox.curselection()  # Получаем текущий выбранный элемент в listbox
        if selection:
            index = selection[0]
            weight_to_remove = self.weights.pop(index)  # Удалить вес из списка
            self.total_weight -= weight_to_remove
            self.total_cost -= self.calculate_cost(weight_to_remove)
            self.weights_listbox.delete(index)  # Удалить элемент из listbox
            self.total_parcels -= 1
        else:
            messagebox.showinfo("Информация", "Выберите вес, который нужно удалить.")
        # Фокусировка после изменения
        self.weight_entry.focus_set()

# Подсчет посылок
    def open_parcels_window(self):
        self.parcels_window = tk.Toplevel(self)
        self.parcels_window.title("Подсчет посылок")

        self.parcels_window.geometry("350x500")

        screen_width = self.parcels_window.winfo_screenwidth()
        screen_height = self.parcels_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 350) // 2- 500
        y_coordinate = (screen_height - 500) // 2

        # Устанавливаем положение окна по центру
        self.parcels_window.geometry(f"350x500+{x_coordinate}+{y_coordinate}")

        self.parcels_price_label = tk.Label(self.parcels_window, text="Введите цену посылки без НДС и нажмите Enter:")
        self.parcels_price_label.pack(pady=10)

        self.parcels_price_entry = tk.Entry(self.parcels_window)
        self.parcels_price_entry.pack(pady=10)
        self.parcels_price_entry.focus_set()

        self.parcels_weights_listbox_label = tk.Label(self.parcels_window, text="Список цен посылок:")
        self.parcels_weights_listbox_label.pack()

        self.parcels_weights_listbox = tk.Listbox(self.parcels_window)
        self.parcels_weights_listbox.pack()

        self.delete_selected_parcel_button = tk.Button(
            self.parcels_window, text="Удалить выбранную посылку", command=self.delete_selected_parcel
        )
        self.delete_selected_parcel_button.pack(pady=10)

        # Кнопка "Закончить подсчет" для открытия календаря
        self.finish_button = tk.Button(self.parcels_window, text="Закончить подсчет", command=self.open_calendar_parcels)
        self.finish_button.pack(pady=(5, 10))

        self.parcels_price_entry.bind("<Return>", self.add_parcel_weight)

        by_label = tk.Label(self.parcels_window, text="by.Borzzz", fg="gray")
        by_label.pack(side=tk.BOTTOM, anchor=tk.SE, padx=10, pady=10)

    def add_parcel_weight(self, event=None):
        try:
            # Получаем значение из виджета Entry и заменяем запятую на точку
            price_entry_text = self.parcels_price_entry.get().replace(',', '.')
            price = float(price_entry_text)

            self.parcels_weights_listbox.insert(tk.END, f"{price} руб.")
            self.parcels_price_entry.delete(0, tk.END)
        except ValueError as e:
            messagebox.showerror("Ошибка", "Введите корректную цену.")
        finally:
            self.parcels_price_entry.focus_set()

    def delete_selected_parcel(self):
        selected_index = self.parcels_weights_listbox.curselection()
        if selected_index:
            self.parcels_weights_listbox.delete(selected_index)

    def open_calendar_parcels(self):
        # Создание календаря
        self.calendar_window = tk.Toplevel(self.parcels_window)
        self.calendar_window.title("Выберите дату")
        self.calendar_window.geometry("300x300")

        screen_width = self.calendar_window.winfo_screenwidth()
        screen_height = self.calendar_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2
        y_coordinate = (screen_height - 300) // 2

        # Устанавливаем положение окна по центру
        self.calendar_window.geometry(f"300x300+{x_coordinate}+{y_coordinate}")

        self.cal = Calendar(self.calendar_window, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        self.cal.pack(pady=10)

        # Кнопка для сохранения результатов
        self.save_button = tk.Button(self.calendar_window, text="Сохранить", command=self.calculate_and_save_parcels)
        self.save_button.pack(pady=10)

    def calculate_and_save_parcels(self):
        try:
            selected_date = self.cal.get_date()
            selected_date = re.sub(r'[\s\\\/.,]', '.', selected_date)
            current_date = datetime.strptime(selected_date, "%d.%m.%Y").date()

            total_parcels = self.parcels_weights_listbox.size()
            total_cost = sum(float(self.parcels_weights_listbox.get(i).split()[0]) for i in range(total_parcels))
            total_cost_with_vat = total_cost * self.NDS  # Учет НДС (20%)

            result_string = (
                f"Итого за {current_date.strftime('%d.%m.%Y')} отправлено {total_parcels} посылок на общую сумму "
                f"{total_cost:.2f} рублей (без НДС) и {total_cost_with_vat:.2f} рублей (с НДС).\n"
            )

            custom_path = self.custom_path
            filename = os.path.join(custom_path, f"Списки посылок.txt")

            with open(filename, "a", encoding='utf-8') as file:
                file.write(result_string)

            messagebox.showinfo("Успешно", "Результаты сохранены.")
            self.parcels_window.destroy()

        except ValueError:
            messagebox.showerror("Ошибка", "Введите дату в правильном формате (дд.мм.гггг).")

# Функция для создания диалогового окна по общему подсчету и ввода даты
    def ask_month_input(self):

        self.month_window = tk.Toplevel(self)
        self.month_window.title("Выберите дату для отчета")
        self.month_window.geometry("400x320")

        screen_width = self.month_window.winfo_screenwidth()
        screen_height = self.month_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 400) // 2 + 500
        y_coordinate = (screen_height - 320) // 2

        # Устанавливаем положение окна по центру
        self.month_window.geometry(f"400x320+{x_coordinate}+{y_coordinate}")

        label_text = "День обязателен, но не учитывается"
        label = tk.Label(self.month_window, text=label_text)
        label.pack(pady=(10, 0))
        
        self.month_calendar = Calendar(self.month_window, selectmode='day', year=2024, month=2, date_pattern='dd.mm.yyyy')

        self.month_calendar.pack(pady=10)

        self.ok_button = tk.Button(self.month_window, text="OK", command=self.get_selected_month)
        self.ok_button.pack(pady=10)

        self.month_window.geometry(f"400x320+{x_coordinate}+{y_coordinate}")

        by_label = tk.Label(self.month_window, text="by.Borzzz", fg="gray")
        by_label.pack(side=tk.BOTTOM, anchor=tk.SE, padx=10, pady=10)


    def get_selected_month(self):
        selected_date_str = self.month_calendar.get_date()
        
        if selected_date_str:
            day, month, year = map(int, selected_date_str.split('.'))
            selected_month = f"{month:02d}.{year}"
            self.calculate_total_for_month(selected_month)
            self.month_window.destroy()
        else:
            print("Дата не выбрана.")

# Это сохранение итогов (за определенный месяц)
    def calculate_total_for_month(self, selected_month):
        try:
            custom_path = self.custom_path
            month, year = map(int, selected_month.split("."))
        
            # Инициализация переменных для подсчета итогов
            total_weight = 0
            total_cost = 0.0
            total_parcels = 0
            total_letters_cost = 0.0
            total_registered_letters_cost = 0.0
            total_simple_letters = 0
            total_registered_letters = 0
            total_parcels_package = 0
            total_cost_package = 0
            total_cost_with_vat_package = 0
        
            # Обработка данных из файла бандеролей
            parcels_file_path = os.path.join(custom_path, "Списки бандеролей.txt")
            if not os.path.exists(parcels_file_path):
                raise FileNotFoundError(f"Файл {parcels_file_path} не найден.")
        
            with open(parcels_file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    date_match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', line)
                    if date_match:
                        line_day, line_month, line_year = map(int, date_match.groups())
                        if line_month == month and line_year == year:
                            total_parcels += int(re.search(r'отправлено (\d+) бандеролей', line).group(1))
                            total_weight += float(re.search(r'весом ([\d.]+) грамм', line).group(1))
                            total_cost += float(re.search(r'на сумму ([\d.]+) рублей', line).group(1))
        
            # Обработка данных из файла простых писем
            simple_letters_file_path = os.path.join(custom_path, "Списки простых писем.txt")
            if not os.path.exists(simple_letters_file_path):
                raise FileNotFoundError(f"Файл {simple_letters_file_path} не найден.")
        
            with open(simple_letters_file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    if selected_month in line:
                        total_simple_letters += int(re.search(r'Количество писем: (\d+)', line).group(1))
                        total_letters_cost += float(re.search(r'Итого: ([\d.]+) руб.', line).group(1))
        
            # Обработка данных из файла заказных писем
            registered_letters_file_path = os.path.join(custom_path, "Списки заказных писем.txt")
            if not os.path.exists(registered_letters_file_path):
                raise FileNotFoundError(f"Файл {registered_letters_file_path} не найден.")
        
            with open(registered_letters_file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    if selected_month in line:
                        total_registered_letters += int(re.search(r'Количество писем: (\d+)', line).group(1))
                        total_registered_letters_cost += float(re.search(r'Итого: ([\d.]+) руб.', line).group(1))

            parcels_file_path = os.path.join(custom_path, "Списки посылок.txt")
            if not os.path.exists(parcels_file_path):
                raise FileNotFoundError(f"Файл {parcels_file_path} не найден.")

            with open(parcels_file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    if selected_month in line:
                        match_parcels = re.search(r'отправлено (\d+) посылок', line)
                        if match_parcels:
                            total_parcels_package += int(match_parcels.group(1))
                        match_cost = re.search(r'на общую сумму ([\d.]+) рублей \(без НДС\)', line)
                        if match_cost:
                            total_cost_package += float(match_cost.group(1))
                        match_cost_with_vat = re.search(r'и ([\d.]+) рублей \(с НДС\)', line)
                        if match_cost_with_vat:
                            total_cost_with_vat_package += float(match_cost_with_vat.group(1))

        
            # Строка результата расчета
            result_string = (f"Итого за {selected_month}:\n"
                             f"Отправлено {total_parcels} бандеролей "
                             f"весом {total_weight} грамм на сумму {total_cost} рублей.\n"
                             f"Отправлено простых писем: {total_simple_letters} на сумму {total_letters_cost} рублей.\n"
                             f"Отправлено заказных писем: {total_registered_letters} на сумму {total_registered_letters_cost} рублей.\n"
                             f"Отправлено посылок: {total_parcels_package} на сумму {total_cost_package} (без НДС) и {total_cost_with_vat_package} (с НДС)\n")
        
            # Показываем результат в message box
            messagebox.showinfo("Итоги за месяц", result_string)
        
            # Сохранение в файл
            output_file_path = os.path.join(custom_path, f"Отчет за {selected_month}.txt")
            with open(output_file_path, 'w', encoding='utf-8') as output_file:
                output_file.write(result_string)
    
        except ValueError as ve:
            messagebox.showerror("Ошибка", str(ve))
        except FileNotFoundError as fe:
            messagebox.showerror("Ошибка", str(fe))

# Кнопка создания обложки
    def open_cover_window(self):
        self.cover_window = tk.Toplevel(self)
        self.cover_window.title("Выберите тип обложки")

        # Размер окна
        window_width = 400
        window_height = 180

        # Получаем размеры экрана
        screen_width = self.cover_window.winfo_screenwidth()
        screen_height = self.cover_window.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # Устанавливаем позицию окна
        self.cover_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # Кнопки
        self.parcel_button = tk.Button(self.cover_window, text="   Обложка почты    ", command=self.open_post)
        self.parcel_button.pack(pady=(20, 0))  # Устанавливаем отступ сверху для первой кнопки

        tk.Label(self.cover_window, text="").pack()

        self.mail_button = tk.Button(self.cover_window, text=" Обложка посылок ", command=self.open_pacage)
        self.mail_button.pack(pady=(10, 20),)  # Устанавливаем отступ снизу для второй кнопки

        # Надпись "by.Borzzz" в нижнем правом углу
        by_label = tk.Label(self.cover_window, text="by.Borzzz", fg="gray")
        by_label.pack(side=tk.BOTTOM, anchor=tk.SE, padx=10, pady=10)


# Обложка на почту
    def open_post(self):
        # Создаем окно
        top = tk.Toplevel(self)

        # Получаем размеры экрана
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - top.winfo_reqwidth()) / 2
        y = (screen_height - top.winfo_reqheight()) / 2

        # Устанавливаем позицию окна
        top.geometry("+%d+%d" % (x, y))

        # Функция для выбора даты
        def get_date():
            selected_date_str = cal.get_date()
            selected_date = datetime.strptime(selected_date_str, "%d.%m.%Y")
            top.destroy()
            self.create_document(selected_date)

        # Календарь для выбора даты
        cal = Calendar(top, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        cal.pack(padx=10, pady=10)

        # Кнопка для выбора даты
        btn_ok = tk.Button(top, text="Выбрать", command=get_date)
        btn_ok.pack(pady=5)

    def create_document(self, selected_date):
        # Создаем документ Word с выбранной датой
        document = DocxDocument()
        
        # Установка шрифта Times New Roman
        run_font = document.styles['Normal'].font
        run_font.name = 'Times New Roman'

        # Добавляем каждую надпись на отдельной строке с выравниванием по центру
        for text in [" ", "Реестры","передачи", "почтовых", "отправлений", f"{selected_date.strftime('%d.%m.%Y')}"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(72)
            run.font.bold = True  # Установка жирного шрифта
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_after = Inches(0.2)  # Пространство после каждой строки
            run.font.name = 'Times New Roman'  # Установка шрифта Times New Roman

        # Устанавливаем ориентацию страницы на книжную
        section = document.sections[0]
        section.orientation = WD_ORIENTATION.PORTRAIT

        # Открываем документ
        temp_file_path = tempfile.mktemp(suffix='.docx')
        document.save(temp_file_path)
        os.startfile(temp_file_path)

# Обложка на посылки
    def open_pacage(self):
        # Создаем окно
        top = tk.Toplevel(self)

        # Получаем размеры экрана
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - top.winfo_reqwidth()) / 2
        y = (screen_height - top.winfo_reqheight()) / 2

        # Устанавливаем позицию окна
        top.geometry("+%d+%d" % (x, y))

        # Функция для выбора даты
        def get_date():
            selected_date_str = cal.get_date()
            selected_date = datetime.strptime(selected_date_str, "%d.%m.%Y")
            top.destroy()
            self.create_document_2(selected_date)

        # Календарь для выбора даты
        cal = Calendar(top, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        cal.pack(padx=10, pady=10)

        # Кнопка для выбора даты
        btn_ok = tk.Button(top, text="Выбрать", command=get_date)
        btn_ok.pack(pady=5)

    def create_document_2(self, selected_date):
        # Создаем документ Word с выбранной датой
        document = DocxDocument()
    
        # Установка шрифта Times New Roman
        run_font = document.styles['Normal'].font
        run_font.name = 'Times New Roman'

        # Добавляем каждую надпись на отдельной строке с выравниванием по центру
        for text in [" ", " ", "Посылки", "за " f"{selected_date.strftime('%m.%Y')}"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(72)
            run.font.bold = True  # Установка жирного шрифта
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_after = Inches(0.2)  # Пространство после каждой строки
            run.font.name = 'Times New Roman'  # Установка шрифта Times New Roman

        section = document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # Открываем документ
        temp_file_path = tempfile.mktemp(suffix='.docx')
        document.save(temp_file_path)
        os.startfile(temp_file_path) 

# Открывается кнопка Настроек
    def open_settings_window(self):
        self.settings_window = tk.Toplevel(self)
        self.settings_window.title("Настройки")
        self.settings_window.geometry("400x710")
        self.settings_window.attributes('-topmost', 'true')

        screen_width = self.settings_window.winfo_screenwidth()
        screen_height = self.settings_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 400) // 2 - 500
        y_coordinate = (screen_height - 710) // 2

        # Устанавливаем положение окна по центру
        self.settings_window.geometry(f"400x710+{x_coordinate}+{y_coordinate}")

        self.settings_title_label = tk.Label(self.settings_window, text="Настройки почтовых тарифов")
        self.settings_title_label.pack(pady=(10, 0))  # небольшой отступ сверху

        # Линия для отделения надписи от остальных элементов
        self.settings_separator = tk.Frame(self.settings_window, height=2, bg="grey")  # создаем рамку с высотой 2 пикселя и цветом серого
        self.settings_separator.pack(fill='x', pady=(5, 10)) 
        
        self.settings_window.attributes('-topmost', 'true')
        
        self.base_cost_label = tk.Label(self.settings_window, text="Стоимость бандероли в 120 грамм:")
        self.base_cost_label.pack(pady=10)
        
        self.base_cost_entry = tk.Entry(self.settings_window)
        self.base_cost_entry.pack(pady=10)
        self.base_cost_entry.insert(0, str(self.BASE_COST))
        
        self.step_cost_label = tk.Label(self.settings_window, text="Стоимость за шаг в 20 грамм:")
        self.step_cost_label.pack(pady=10)
        
        self.step_cost_entry = tk.Entry(self.settings_window)
        self.step_cost_entry.pack(pady=10)
        self.step_cost_entry.insert(0, str(self.STEP_COST))
        
        self.letter_cost_label = tk.Label(self.settings_window, text="Стоимость простого письма:")
        self.letter_cost_label.pack(pady=10)
        
        self.letter_cost_entry = tk.Entry(self.settings_window)
        self.letter_cost_entry.pack(pady=10)
        self.letter_cost_entry.insert(0, str(self.LETTER_COST))

        self.registered_letter_cost_label = tk.Label(self.settings_window, text="Стоимость заказного письма:")
        self.registered_letter_cost_label.pack(pady=10)
        
        self.registered_letter_cost_entry = tk.Entry(self.settings_window)
        self.registered_letter_cost_entry.pack(pady=10)
        self.registered_letter_cost_entry.insert(0, str(self.REGISTERED_LETTER_COST))

        self.nds_entry_label = tk.Label(self.settings_window, text="Переводим надбавку НДС:")
        self.nds_entry_label.pack(pady=10)

        self.nds_entry = tk.Entry(self.settings_window)
        self.nds_entry.pack(pady=10)
        self.nds_entry.insert(0, str(self.NDS))

        self.settings_separator = tk.Frame(self.settings_window, height=2, bg="grey")  # создаем рамку с высотой 2 пикселя и цветом серого
        self.settings_separator.pack(fill='x', pady=(10, 10))

        self.settings_title_label = tk.Label(self.settings_window, text="Настройка пути хранения файлов.\n Без необходимости не трогать!!!")
        self.settings_title_label.pack(pady=(10, 0))  # небольшой отступ сверху

        self.custom_path_entry = tk.Entry(self.settings_window)
        self.custom_path_entry.pack(pady=10)
        self.custom_path_entry.insert(0, str(self.custom_path))

        self.custom_path_button = tk.Button(self.settings_window, text="Выбрать путь", command=self.select_custom_path)
        self.custom_path_button.pack(pady=10)

        self.settings_separator = tk.Frame(self.settings_window, height=2, bg="grey")  # создаем рамку с высотой 2 пикселя и цветом серого
        self.settings_separator.pack(fill='x', pady=(10, 10))
        
        self.save_settings_button = tk.Button(
            self.settings_window, text="Сохранить", command=self.save_settings
        )
        self.save_settings_button.pack(pady=10)

        by_label = tk.Label(self.settings_window, text="by.Borzzz", fg="gray")
        by_label.pack(side=tk.BOTTOM, anchor=tk.SE, padx=10, pady=10)

# Окно выбора пути сохранения
    def select_custom_path(self):
        path = askdirectory()  # Показать диалоговое окно и вернуть выбранный путь
        if path:
            self.custom_path_entry.delete(0, tk.END)  # Очистка текущего содержимого Entry
            self.custom_path_entry.insert(0, path)  # Вставить выбранный путь

# Сохраняем настройки
    def save_settings(self):
        try:
            self.BASE_COST = float(self.base_cost_entry.get())
            self.STEP_COST = float(self.step_cost_entry.get())
            self.LETTER_COST = float(self.letter_cost_entry.get())
            self.NDS = float(self.nds_entry.get())
            self.custom_path = self.custom_path_entry.get()
            self.save_settings_to_file()  # вызов метода для сохранения всех настроек, включая путь
            self.settings_window.destroy()
            messagebox.showinfo("Успех", "Настройки сохранены.")
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректные числовые значения.")

# Это сохранение настроек, чтобы они не сбивались при закрытии    
    def save_settings_to_file(self):
        settings = {
            'BASE_COST': self.BASE_COST,
            'STEP_COST': self.STEP_COST,
            'LETTER_COST': self.LETTER_COST,
            'REGISTERED_LETTER_COST' : self.REGISTERED_LETTER_COST,
            'NDS' : self.NDS,
            'CUSTOM_PATH': self.custom_path
        }
        settings_path = os.path.join(os.getenv('APPDATA'), 'settings.json')
        os.makedirs(os.path.dirname(settings_path), exist_ok=True)  # Создаем директорию, если она не существует
        with open(settings_path, 'w') as f:
            json.dump(settings, f)

        print(f'Файл настроек сохранен по пути: {settings_path}')

# Отсюда загружаем настройки приложения
    def load_settings_from_file(self):
        settings_path = os.path.join(os.getenv('APPDATA'), 'settings.json')
        try:
            with open(settings_path, 'r') as f:
                settings = json.load(f)
            self.BASE_COST = settings.get('BASE_COST', self.BASE_COST)
            self.STEP_COST = settings.get('STEP_COST', self.STEP_COST)
            self.LETTER_COST = settings.get('LETTER_COST', self.LETTER_COST)
            self.REGISTERED_LETTER_COST = settings.get('REGISTERED_LETTER_COST', self.REGISTERED_LETTER_COST)
            self.NDS = settings.get('NDS', self.NDS)
            self.custom_path = settings.get('CUSTOM_PATH', os.path.expanduser('~'))
        except FileNotFoundError:
            self.custom_path = os.path.expanduser('~') # Файл с настройками отсутствует, будут использованы значения по умолчанию
            self.save_settings_to_file()  # Создаем файл настроек со значениями по умолчанию
        except json.JSONDecodeError:
            messagebox.showerror("Ошибка", "Ошибка чтения настроек. Проверьте файл настроек.")





if __name__ == "__main__":
    app = App()
    app.mainloop()